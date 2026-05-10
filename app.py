"""Meeting Minutes Summarizer — Phase 1: minimal end-to-end loop.

Paste Korean meeting text, click 요약, get a .docx summary in output/.
"""
from datetime import datetime
from pathlib import Path
import tkinter as tk
from tkinter import scrolledtext

from openai import OpenAI
from docx import Document

MODEL = "gpt-4o-mini"
OUTPUT_DIR = Path(__file__).parent / "output"


def summarize_text(meeting_text: str, api_key: str) -> str:
    """Call gpt-4o-mini to summarize Korean meeting minutes."""
    client = OpenAI(api_key=api_key.strip(), timeout=60)
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "system", "content": "당신은 한국어 회의록 요약 도우미입니다."},
            {"role": "user",   "content": f"다음 한국어 회의록을 한국어로 간결하게 요약해 주세요:\n\n{meeting_text}"},
        ],
    )
    return response.choices[0].message.content


def write_docx(summary: str) -> Path:
    """Write summary to output/summary-{YYYYMMDD-HHMMSS}.docx. Returns the path."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    path = OUTPUT_DIR / f"summary-{timestamp}.docx"
    doc = Document()
    doc.add_heading("회의 요약", level=1)
    doc.add_paragraph(summary)
    doc.save(path)
    return path


def main() -> None:
    """Build Tkinter UI, wire on_click handler, run mainloop."""
    root = tk.Tk()
    root.title("회의록 요약기")
    root.geometry("700x600")

    tk.Label(root, text="OpenAI API 키:").pack(anchor="w", padx=8, pady=(8, 0))
    api_key_entry = tk.Entry(root, show="*")
    api_key_entry.pack(fill="x", padx=8)

    tk.Label(root, text="회의록 붙여넣기:").pack(anchor="w", padx=8, pady=(8, 0))
    input_text = scrolledtext.ScrolledText(root, height=15, wrap="word")
    input_text.pack(fill="both", expand=True, padx=8)

    summarize_btn = tk.Button(root, text="요약")
    summarize_btn.pack(pady=8)

    status_var = tk.StringVar(value="대기 중")
    tk.Label(root, textvariable=status_var, fg="gray").pack(anchor="w", padx=8)

    tk.Label(root, text="요약 결과:").pack(anchor="w", padx=8, pady=(8, 0))
    result_text = scrolledtext.ScrolledText(root, height=10, wrap="word")
    result_text.pack(fill="both", expand=True, padx=8, pady=(0, 8))

    def on_click():
        api_key = api_key_entry.get().strip()
        text = input_text.get("1.0", "end-1c").strip()
        if not api_key or not text:
            status_var.set("API 키와 회의록을 모두 입력해주세요.")
            return
        summarize_btn.config(state="disabled")
        status_var.set("요약 중... (창이 잠시 멈출 수 있습니다)")
        # NOTE: sync API call freezes UI for ~5-15s. Threading deferred to Phase 5 (UI-03).
        root.update_idletasks()
        try:
            summary = summarize_text(text, api_key)
            path = write_docx(summary)
            result_text.delete("1.0", "end")
            result_text.insert("1.0", summary)
            status_var.set(f"완료. 저장됨: {path}")
        except Exception as e:
            status_var.set(f"에러: {e}")
        finally:
            summarize_btn.config(state="normal")

    summarize_btn.config(command=on_click)
    root.mainloop()


if __name__ == "__main__":
    main()
